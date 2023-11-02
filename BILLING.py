import sqlite3
import pandas as pd
import tkinter as tk
import matplotlib.pyplot as plt
import tkinter.messagebox
import openpyxl
from PATDELSU import P_display_res
from docx import Document
conn=sqlite3.connect("MDBA.db")

#variables
rootB=tk.Tk()
rootB.withdraw()
patient_id_entry = tk.Entry(rootB)
patient_id_entry.place(x=200, y=150)
patient_id_entry.pack()
 
import matplotlib.pyplot as plt
patient_id_entry = tkinter.Entry(rootB)
patient_id_entry.pack()

def plot_data(patient_id):
    
    # Execute a SQL query to retrieve the data
    cur = conn.cursor()
    
    p=list(cur.execute('select * from employee where PATIENT_ID=?',(patient_id,)))
    cur.execute("SELECT MERA,AGE,muoar FROM employee WHERE PATIENT_ID=?", (patient_id,))
    if (len(p)==0):
        errorS=tkinter.Label(rootB,text="PATIENT RECORD NOT FOUND")
        errorS.place(x='120',y="300")
    else:
        # Store the data in variables
        x = []
        y = []
        z = []
        for row in cur:
            x.append(row[0])
            y.append(row[1])
            z.append(row[2])

    # Close the cursor and connection
    cur.close()
    plt.xlabel("DATE")
    plt.ylabel("MΥΩΠΙΑ")
    plt.title("ΕΞΕΛΙΞΗ ΜΥΩΠΙΑΣ")

    # Create the plot
    plt.plot(x, y, label = "ΔΕΞΙ ΜΑΤΙ")
    plt.plot(x, z,label = "ΑΡΙΣΤΕΡΟ ΜΑΤΙ")
    plt.legend()
    plt.show()

def plot_data_presv(patient_id):
    
    # Execute a SQL query to retrieve the data
    cur = conn.cursor()
    
    p=list(cur.execute('select * from employee where PATIENT_ID=?',(patient_id,)))
    cur.execute("SELECT MERA,SEX,presvar FROM employee WHERE PATIENT_ID=?", (patient_id,))
    if (len(p)==0):
        errorS=tkinter.Label(rootB,text="PATIENT RECORD NOT FOUND")
        errorS.place(x='120',y="300")
    else:
        # Store the data in variables
        x = []
        y = []
        z = []
        for row in cur:
            x.append(row[0])
            y.append(row[1])
            z.append(row[2])

    # Close the cursor and connection
    cur.close()
    plt.xlabel("DATE")
    plt.ylabel("ΠΡΕΣΒΥΩΠΙΑ")
    plt.title("ΕΞΕΛΙΞΗ ΠΡΕΣΒΥΩΠΙΑΣ")

    # Create the plot
    plt.plot(x, y, label = "ΔΕΞΙ ΜΑΤΙ")
    plt.plot(x, z,label = "ΑΡΙΣΤΕΡΟ ΜΑΤΙ")
    plt.legend()
    plt.show()

def plot_data_uper(patient_id):
    
    # Execute a SQL query to retrieve the data
    cur = conn.cursor()
    
    p=list(cur.execute('select * from employee where PATIENT_ID=?',(patient_id,)))
    cur.execute("SELECT MERA,SAL,uperari FROM employee WHERE PATIENT_ID=?", (patient_id,))
    if (len(p)==0):
        errorS=tkinter.Label(rootB,text="PATIENT RECORD NOT FOUND")
        errorS.place(x='120',y="300")
    else:
        # Store the data in variables
        x = []
        y = []
        z = []
        for row in cur:
            x.append(row[0])
            y.append(row[1])
            z.append(row[2])

    # Close the cursor and connection
    cur.close()
    plt.xlabel("DATE")
    plt.ylabel("υπερμετρωπία")
    plt.title("ΕΞΕΛΙΞΗ υπερμετρωπίας")

    # Create the plot
    plt.plot(x, y, label = "ΔΕΞΙ ΜΑΤΙ")
    plt.plot(x, z,label = "ΑΡΙΣΤΕΡΟ ΜΑΤΙ")
    plt.legend()
    plt.show()

def plot_data_astigm(patient_id):
    
    # Execute a SQL query to retrieve the data
    cur = conn.cursor()
    
    p=list(cur.execute('select * from employee where PATIENT_ID=?',(patient_id,)))
    cur.execute("SELECT MERA,EXP,astar FROM employee WHERE PATIENT_ID=?", (patient_id,))
    if (len(p)==0):
        errorS=tkinter.Label(rootB,text="PATIENT RECORD NOT FOUND")
        errorS.place(x='120',y="300")
    else:
        # Store the data in variables
        x = []
        y = []
        z = []
        for row in cur:
            x.append(row[0])
            y.append(row[1])
            z.append(row[2])

    # Close the cursor and connection
    cur.close()
    plt.xlabel("DATE")
    plt.ylabel("AΣΤΙΓΜΑΤΙΣΜΟΣ")
    plt.title("ΕΞΕΛΙΞΗ AΣΤΙΓΜΑΤΙΣΜΟΥ")

    # Create the plot
    plt.plot(x, y, label = "ΔΕΞΙ ΜΑΤΙ")
    plt.plot(x, z,label = "ΑΡΙΣΤΕΡΟ ΜΑΤΙ")
    plt.legend()
    plt.show()




L1=None
L2=None
L3=None
L4=None
 

def exitt():
    rootB.destroy()


def show_patients_table():
    # Get the data from the database
    cursor = conn.execute('SELECT * FROM PATIENT')
    data = []
    for row in cursor:
        data.append(row)

    # Create a dataframe from the data
    df = pd.DataFrame(data, columns=['AMKA', 'ONOMA', 'ASFALIA'])

    # Save the DataFrame to an Excel file
    df.to_excel('asthenis.xlsx', index=False, encoding='utf-8')

    # Open the Excel file in Word
    document = openpyxl.load_workbook('asthenis.xlsx')
    document.save('asthenis.docx')

    # Display the dataframe
    print(df)

    top = tkinter.Toplevel()
    top.title("Message")
    tkinter.Label(top, text="ΤΟ ΑΡΧΕΙΟ ΕΚΤΥΠΩΘΗΚΕ ΣΕ ΑΡΧΕΙΟ").pack()
    tkinter.Button(top, text="Close", command=top.destroy).pack()

 
def show_patients_FARMAKA():
    # Get the data from the database
    cursor = conn.execute('SELECT * FROM ROOM')
    data = []
    for row in cursor:
        data.append(row)

    # Create a dataframe from the data
    df = pd.DataFrame(data, columns=['AMKA', 'ΙD ΡΑΝΤΕΒΟΥ', 'ΤΥΠΟΣ ΦΑΡΜΑΚΩΝ','ΣΥΜΠΤΩΜΑΤΟΛΟΓΙΑ','ΗΜΕΡΟΜΙΝΙΑ','ΑΠΟΤΕΛΕΣΜΑΤΑ'])

    # Save the DataFrame to an Excel file
    df.to_excel('FARMAKA.xlsx', index=False, encoding='utf-8')

    # Open the Excel file in Word
    document = openpyxl.load_workbook('FARMAKA.xlsx')
    document.save('FARMAKA.docx')

    # Display the dataframe
    print(df)

    top = tkinter.Toplevel()
    top.title("Message")
    tkinter.Label(top, text="ΤΟ ΑΡΧΕΙΟ ΕΚΤΥΠΩΘΗΚΕ ΣΕ ΑΡΧΕΙΟ").pack()
    tkinter.Button(top, text="Close", command=top.destroy).pack()

def show_patients_EKSETA():
    # Get the data from the database
    cursor = conn.execute('SELECT * FROM employee')
    data = []
    for row in cursor:
        data.append(row)

    # Create a dataframe from the data
    df = pd.DataFrame(data, columns=['AMKA', 'ONOMA', 'ΠΡΕΣΒΥΩΠΙΑ ΔΕΞΙ','MΥΩΠΙΑ ΔΕΞΙ','ΙΔ ΡΑΝΤΕΒΟΥ','υπερμετρωπία ΔΕΞΙ','AΣΤΙΓΜΑΤΙΣΜΟΣ δεξι','ΚΟΣΤΟΣ','PHONE','idapo','ΗΜΕΡΟΜΗΝΙΑ','Βαθμός μυωπίας αριστερό μάτι','βαθμός υπερμετρωπίας αριστερό μάτι','Βαθμος πρεσβυωπία αριστερό μάτι','Πίεση αριστερού ματιού','βαθμός αστιγματισμού αριστερό'])
                      
    # Save the DataFrame to an Excel file
    df.to_excel('ekseta.xlsx', index=False, encoding='utf-8')

    # Open the Excel file in Word
    document = openpyxl.load_workbook('ekseta.xlsx')
    document.save('ekseta.docx')

    # Display the dataframe
    print(df)

    top = tkinter.Toplevel()
    top.title("Message")
    tkinter.Label(top, text="ΤΟ ΑΡΧΕΙΟ ΕΚΤΥΠΩΘΗΚΕ ΣΕ ΑΡΧΕΙΟ").pack()
    tkinter.Button(top, text="Close", command=top.destroy).pack()
   
def create_patient_report(df):
    # Create a new Word document
    document = Document()

    # Add the DataFrame as a table to the document
    document.add_table(df.shape[0]+1, df.shape[1])

    # Add the column names to the first row of the table
    for j in range(df.shape[-1]):
        document.tables[0].cell(0,j).text = df.columns[j]

    # Add the data to the rest of the table
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            document.tables[0].cell(i+1,j).text = str(df.values[i,j])

    # Save the document
    document.save("patient_data.docx")


 

label = tkinter.Label(rootB, text="Enter patient ID:")
entry_field = tkinter.Entry(rootB)
    

def BILLING():
    global rootB,L1,treat1,P_id,dd,cost,med,med_q,price,treat_1,treat_2,cost_t,j,jj,jjj,jjjj,L2,L3,L4
    rootB=tk.Tk()
    rootB.geometry("450x350")
    rootB.title("ΠΡΟΒΟΛΗ ΑΠΟΤΕΛΕΣΜΑΤΩΝ")
    head=tk.Label(rootB,text="ΟΛΑ ΤΑ ΔΕΔΟΜΕΝΑ",font="Arial 16 bold italic",fg='grey')
    head.place(x=100,y=10)

    # Create the Entry widget
     # Create the Entry widget
    patient_id_entry = tk.Entry(rootB)
    patient_id_entry.place(x='20',y='190')
    l11 = tkinter.Label(rootB, text="ΒΑΛΕ ΑΜΚΑ ΓΙΑ ΔΙΑΓΡΑΜΜΑ")
    l11.place(x=20,y=170)
    button = tk.Button(rootB, text='ΜΥΩΠΙΑΣ', command=lambda: plot_data(patient_id_entry.get()))
    button.place(x='35',y='215')

    button = tk.Button(rootB, text='ΠΡΕΣΒΥΩΠΙΑΣ', command=lambda: plot_data_presv(patient_id_entry.get()))
    button.place(x='130',y='215')

    button = tk.Button(rootB, text='ΥΠΕΡΜΕΤΡΩΠΙΑ', command=lambda: plot_data_uper(patient_id_entry.get()))
    button.place(x='245',y='215')

    button = tk.Button(rootB, text='ΑΣΤΙΓΜΑΤΙΣΜΟΥ', command=lambda: plot_data_astigm(patient_id_entry.get()))
    button.place(x='360',y='215')

    ee=tk.Button(rootB,text="EXIT",command=exitt)
    ee.place(x='310',y='310')
    show_patients_table_button = tk.Button(rootB, text='ΟΛΟΙ ΟΙ ΑΣΘΕΝΕΙΣ', command=show_patients_table)
    show_patients_table_button.place(x='20',y='100')

    show_patients_table_button = tk.Button(rootB, text='ΟΛΕΣ ΟΙ ΦΑΡΜΑΚΕΥΤΙΚΕΣ ΑΓΩΓΕΣ', command=show_patients_FARMAKA)
    show_patients_table_button.place(x='150',y='100')
    
    show_patients_table_button = tk.Button(rootB, text='ΟΛΕΣ ΟΙ ΕΞΕΤΑΣΕΙΣ', command=show_patients_EKSETA)
    show_patients_table_button.place(x='20',y='130')
     
    SEARCH=tk.Button(rootB,text="  ΟΛΑ ΤΑ ΡΑΝΤΕΒΟΥ ΑΣΘΕΝΗ  ",command=P_display_res)
    SEARCH.place(x='20',y='60')
    rootB.mainloop()